from io import BytesIO
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from googlesearch import search
import google.generativeai as genai
import re

# Configure Google Gemini API Key
GEMINI_API_KEY = "AIzaSyBKeiRT5dnkvCjIsuSYygKCkK5qUDbcOFQ"

# Initialize Gemini API
genai.configure(api_key=GEMINI_API_KEY)

def call_gemini_api(prompt):
    """Calls the Gemini API to process text."""
    model = genai.GenerativeModel(model_name="gemini-pro")
    response = model.generate_content(prompt)
    return response.text.strip() if response else None

def scrape_web_content(url):
    """Scrapes web content from the given URL."""
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        
        paragraphs = soup.find_all('p')
        content = ' '.join([para.text for para in paragraphs[:5]])  # First 5 paragraphs

        return clean_text(content) if content else "Content not available"
    except requests.exceptions.RequestException as e:
        print(f"Error scraping {url}: {e}")
        return "Content not available"

def clean_text(text):
    """Removes unnecessary characters and extra spaces from text."""
    text = re.sub(r'\*+', '', text)  # Remove asterisks
    text = re.sub(r'\s+', ' ', text).strip()  # Remove multiple spaces
    return text

def search_google_free(topic, num_results=3):
    """Fetch top Google search links using a free method."""
    return [url for url in search(topic, num_results=num_results)]

def add_page_border(doc):
    """Adds a border to all pages of the document."""
    sect = doc.sections[0]
    border = OxmlElement("w:pgBorders")
    border.set(qn("w:offsetFrom"), "page")

    for side in ["top", "left", "bottom", "right"]:
        border_tag = OxmlElement(f"w:{side}")
        border_tag.set(qn("w:val"), "single")
        border_tag.set(qn("w:sz"), "12")  # Border thickness
        border_tag.set(qn("w:space"), "4")
        border_tag.set(qn("w:color"), "000000")  # Black border
        border.append(border_tag)

    sect._sectPr.append(border)

def add_headers_and_footers(doc, topic):
    """Adds a topic in the header and page numbers in the footer."""
    sect = doc.sections[0]

    # Header
    header = sect.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.text = topic
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(14)
    run.bold = True

    # Footer with Page Numbers
    footer = sect.footer
    footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    page_num = OxmlElement("w:fldSimple")
    page_num.set(qn("w:instr"), "PAGE")
    footer_paragraph._element.append(page_num)

def identify_and_bold_subtopics(text, doc):
    """Uses Gemini to extract subtopics and bold them in the text."""
    subtopics_prompt = f"Extract key subtopics from the following text and return them as a comma-separated list:\n{text}"
    subtopics = call_gemini_api(subtopics_prompt)

    if subtopics:
        subtopics_list = [s.strip() for s in subtopics.split(",")]
    else:
        subtopics_list = []

    for line in text.split("\n"):
        para = doc.add_paragraph()
        words = line.split(" ")
        for word in words:
            run = para.add_run(word + " ")
            if word.strip() in subtopics_list:
                run.bold = True

def generate_notes(transcript, topic):
    """Generates notes, refines them with Gemini, and saves to a Word document."""
    doc = Document()
    add_page_border(doc)  # Add border to pages
    add_headers_and_footers(doc, topic)  # Add header and footer

    doc.add_heading(f"Lecture Notes on {topic}", level=1)

    # 1. Transcript Summary
    doc.add_heading("1. Transcript Summary", level=2)
    refined_transcript = call_gemini_api(f"Refine and summarize the following transcript:\n{transcript}")
    identify_and_bold_subtopics(clean_text(refined_transcript) if refined_transcript else clean_text(transcript), doc)

    # 2. Wikipedia Summary
    doc.add_heading("2. Wikipedia Summary", level=2)
    wiki_url = f"https://en.wikipedia.org/wiki/{topic.replace(' ', '_')}"
    wiki_content = scrape_web_content(wiki_url)
    refined_wiki_content = call_gemini_api(f"Refine and summarize the following Wikipedia content:\n{wiki_content}")
    identify_and_bold_subtopics(clean_text(refined_wiki_content) if refined_wiki_content else clean_text(wiki_content), doc)

    # 3. Additional Web Content
    doc.add_heading("3. Additional Web Content", level=2)
    related_links = search_google_free(topic, num_results=3)
    for link in related_links:
        web_content = scrape_web_content(link)
        refined_web_content = call_gemini_api(f"Refine and summarize the following article and add some additional information:\n{web_content}")
        identify_and_bold_subtopics(clean_text(refined_web_content) if refined_web_content else clean_text(web_content), doc)

    # 4. Key Takeaways
    doc.add_heading("4. Key Takeaways", level=2)
    key_points = [
        "Introduction to the topic",
        "Key discoveries and developments",
        "Impact on society",
        "Future advancements and applications"
    ]
    refined_key_points = call_gemini_api(f"Generate detailed bullet points for the following key areas related to {topic}:\n{key_points}")
    identify_and_bold_subtopics(clean_text(refined_key_points) if refined_key_points else '\n'.join(key_points), doc)

    # Sources
    doc.add_heading("Sources", level=2)
    for source in related_links:
        doc.add_paragraph(f"- {source}")

    filename = "Detailed_Lecture_Notes.docx"
    doc.save(filename)
    print(f"✅ Notes saved as {filename}")

# Example usage
transcript_text = "This is a sample transcript about Hyperledger Fabric."
topic_of_interest = "Hyperledger Fabric"
generate_notes(transcript_text, topic_of_interest)
