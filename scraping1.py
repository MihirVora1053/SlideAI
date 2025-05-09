# import requests
# import json
# from bs4 import BeautifulSoup
# from docx import Document
# from io import BytesIO
# from docx.shared import Pt
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx.oxml import OxmlElement
# from docx.oxml.ns import qn
# from googlesearch import search
# import google.generativeai as genai
# import re

# # Configure Google Gemini API Key
# # GEMINI_API_KEY = "AIzaSyBKeiRT5dnkvCjIsuSYygKCkK5qUDbcOFQ"
# GEMINI_API_KEY = "AIzaSyDDiIZWi7KPMfb5e_ur5JhQm87_UXAk48c"

# # Initialize Gemini API
# genai.configure(api_key=GEMINI_API_KEY)

# def call_gemini_api(prompt):
#     """Calls the Gemini API to process text."""
#     model = genai.GenerativeModel(model_name="gemini-pro")
#     response = model.generate_content(prompt)
#     return response.text.strip() if response else None

# def scrape_web_content(url):
#     """Scrapes web content from the given URL."""
#     try:
#         response = requests.get(url, timeout=10)
#         response.raise_for_status()
#         soup = BeautifulSoup(response.text, 'html.parser')

#         paragraphs = soup.find_all('p')
#         content = ' '.join([para.text for para in paragraphs[:5]])  # First 5 paragraphs

#         return clean_text(content) if content else "Content not available"
#     except requests.exceptions.RequestException as e:
#         print(f"Error scraping {url}: {e}")
#         return "Content not available"

# def clean_text(text):
#     """Removes unnecessary characters and extra spaces from text."""
#     text = re.sub(r'\*+', '', text)  # Remove asterisks
#     text = re.sub(r'\s+', ' ', text).strip()  # Remove multiple spaces
#     return text

# def search_google_free(topic, num_results=3):
#     """Fetch top Google search links using a free method."""
#     return [url for url in search(topic, num_results=num_results)]

# def add_page_border(doc):
#     """Adds a border to all pages of the document."""
#     sect = doc.sections[0]
#     border = OxmlElement("w:pgBorders")
#     border.set(qn("w:offsetFrom"), "page")

#     for side in ["top", "left", "bottom", "right"]:
#         border_tag = OxmlElement(f"w:{side}")
#         border_tag.set(qn("w:val"), "single")
#         border_tag.set(qn("w:sz"), "12")  # Border thickness
#         border_tag.set(qn("w:space"), "4")
#         border_tag.set(qn("w:color"), "000000")  # Black border
#         border.append(border_tag)

#     sect._sectPr.append(border)

# def add_headers_and_footers(doc, topic):
#     """Adds a topic in the header and page numbers in the footer."""
#     sect = doc.sections[0]

#     # Header
#     header = sect.header
#     paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
#     paragraph.text = topic
#     paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     run = paragraph.runs[0]
#     run.font.size = Pt(14)
#     run.bold = True

#     # Footer with Page Numbers
#     footer = sect.footer
#     footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
#     footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

#     page_num = OxmlElement("w:fldSimple")
#     page_num.set(qn("w:instr"), "PAGE")
#     footer_paragraph._element.append(page_num)

# def identify_and_bold_subtopics(text, doc):
#     """Uses Gemini to extract subtopics and bold them in the text."""
#     subtopics_prompt = f"Extract key subtopics from the following text and return them as a comma-separated list:\n{text}"
#     subtopics = call_gemini_api(subtopics_prompt)

#     if subtopics:
#         subtopics_list = [s.strip() for s in subtopics.split(",")]
#     else:
#         subtopics_list = []

#     for line in text.split("\n"):
#         para = doc.add_paragraph()
#         words = line.split(" ")
#         for word in words:
#             run = para.add_run(word + " ")
#             if word.strip() in subtopics_list:
#                 run.bold = True


# def generate_notes(transcript, topic):
#     """Generates notes, refines them with Gemini, and saves to a Word document."""
#     doc = Document()
#     add_page_border(doc)  # Add border to pages
#     add_headers_and_footers(doc, topic)  # Add header and footer

#     doc.add_heading(f"Lecture Notes on {topic}", level=1)

#     # 1. Transcript Summary
#     doc.add_heading("1. Transcript Summary", level=2)
#     refined_transcript = call_gemini_api(f"Refine and summarize the following transcript:\n{transcript}")
#     identify_and_bold_subtopics(clean_text(refined_transcript) if refined_transcript else clean_text(transcript), doc)

#     # 2. Wikipedia Summary
#     doc.add_heading("2. Wikipedia Summary", level=2)
#     wiki_url = f"https://en.wikipedia.org/wiki/{topic.replace(' ', '_')}"
#     wiki_content = scrape_web_content(wiki_url)
#     refined_wiki_content = call_gemini_api(f"Refine and summarize the following Wikipedia content:\n{wiki_content}")
#     identify_and_bold_subtopics(clean_text(refined_wiki_content) if refined_wiki_content else clean_text(wiki_content), doc)

#     # 3. Additional Web Content
#     doc.add_heading("3. Additional Web Content", level=2)
#     related_links = search_google_free(topic, num_results=3)
#     for link in related_links:
#         web_content = scrape_web_content(link)
#         refined_web_content = call_gemini_api(f"Refine and summarize the following article and add some additional information:\n{web_content}")
#         identify_and_bold_subtopics(clean_text(refined_web_content) if refined_web_content else clean_text(web_content), doc)

#     # 4. Key Takeaways
#     doc.add_heading("4. Key Takeaways", level=2)
#     key_points = [
#         "Introduction to the topic",
#         "Key discoveries and developments",
#         "Impact on society",
#         "Future advancements and applications"
#     ]
#     refined_key_points = call_gemini_api(f"Generate detailed bullet points for the following key areas related to {topic}:\n{key_points}")
#     identify_and_bold_subtopics(clean_text(refined_key_points) if refined_key_points else '\n'.join(key_points), doc)

#     # Sources
#     doc.add_heading("Sources", level=2)
#     for source in related_links:
#         doc.add_paragraph(f"- {source}")

#     filename = "Detailed_Lecture_Notes.docx"
#     doc.save(filename)
#     print(f"✅ Notes saved as {filename}")

# if __name__ == "__main__":
#     sample_transcript = """Artificial intelligence is transforming industries worldwide.
#     Machine learning models improve decision-making, automate tasks, and enhance efficiency.
#     However, AI ethics and biases remain crucial concerns.
#     The future of AI will be shaped by policy decisions and technological advancements."""

#     print("Refined Transcript:", call_gemini_api(sample_transcript))
#     generate_notes(sample_transcript, "Artificial Intelligence")


import requests
import json
from bs4 import BeautifulSoup
from docx import Document
from nltk.tokenize import sent_tokenize
from io import BytesIO
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from googlesearch import search
import google.generativeai as genai
import re
import os
from dotenv import load_dotenv
from urllib.parse import urljoin

# from spire.doc import *
# from spire.doc.common import *
from PIL import Image


load_dotenv()


# Initialize Gemini API
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))


def call_gemini_api(prompt):
    """Calls the Gemini API to process text."""
    try:
        model = genai.GenerativeModel(model_name="gemini-2.0-flash")
        response = model.generate_content(prompt)
        return response.text.strip() if response and response.text else None
    except Exception as e:
        print(f"Error calling Gemini API: {e}")
        return None


def scrape_web_content(url):
    """Scrapes web content from the given URL."""
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")

        paragraphs = soup.find_all("p")
        content = " ".join([para.text for para in paragraphs[:5]])  # First 5 paragraphs

        return clean_text(content) if content else "Content not available"
    except requests.exceptions.RequestException as e:
        print(f"Error scraping {url}: {e}")
        return "Content not available"


def clean_geeksforgeeks_content(raw_content):
    """
    Cleans raw GFG content by removing markdown headers, asterisks, and excess backticks.
    Also handles spacing for code blocks.
    """
    lines = raw_content.strip().split("\n")
    cleaned_lines = []

    for line in lines:
        # Remove markdown bold (**), headers (###, ####), and bullet markers (*)
        line = re.sub(r"\*{1,2}(.*?)\*{1,2}", r"\1", line)
        line = re.sub(r"^#+\s*", "", line)
        line = re.sub(r"^\*+\s*", "- ", line)

        # Strip leading/trailing whitespace
        stripped = line.strip()

        if stripped == "```python" or stripped == "```":
            cleaned_lines.append("")  # spacing before/after code blocks
        elif not stripped:
            cleaned_lines.append("")
        else:
            cleaned_lines.append(stripped)

    return "\n".join(cleaned_lines)


def clean_text(text):
    """Removes unnecessary characters and extra spaces from text."""
    text = re.sub(r"\*+", "", text)  # Remove asterisks
    text = re.sub(r"\s+", " ", text).strip()  # Remove multiple spaces
    return text


def search_google_free(topic, num_results=3):
    """Fetch top Google search links using a free method."""
    return [url for url in search(topic, num_results=num_results)]


def add_page_border(doc):
    """Adds a border to all pages of the document."""
    sect = doc.sections[0]
    border = OxmlElement("w:pgBorders")
    border.set(qn("w:offsetFrom"), "page")

    for side in ["top", "left", "bottom", "right"]:
        border_tag = OxmlElement(f"w:{side}")
        border_tag.set(qn("w:val"), "single")
        border_tag.set(qn("w:sz"), "12")  # Border thickness
        border_tag.set(qn("w:space"), "4")
        border_tag.set(qn("w:color"), "000000")  # Black border
        border.append(border_tag)

    sect._sectPr.append(border)


def add_headers_and_footers(doc, topic):
    """Adds a topic in the header and page numbers in the footer."""
    sect = doc.sections[0]

    # Header
    header = sect.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.text = topic
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(14)
    run.bold = True

    # Footer with Page Numbers
    footer = sect.footer
    footer_paragraph = (
        footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    )
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    page_num = OxmlElement("w:fldSimple")
    page_num.set(qn("w:instr"), "PAGE")
    footer_paragraph._element.append(page_num)


def identify_and_bold_subtopics(text, doc):
    """Uses Gemini to extract subtopics and bold them in the text."""
    if not text:
        return
    subtopics_prompt = f"Extract key subtopics from the following text and return them as a comma-separated list:\n{text}"
    subtopics = call_gemini_api(subtopics_prompt)

    if subtopics:
        subtopics_list = [s.strip() for s in subtopics.split(",")]
    else:
        subtopics_list = []

    for line in text.split("\n"):
        para = doc.add_paragraph()
        words = line.split(" ")
        for word in words:
            run = para.add_run(word + " ")
            if word.strip() in subtopics_list:
                run.bold = True


def extract_text_and_images(html):
    """
    Parses the HTML to extract text and image URLs.
    Returns (cleaned_text, list_of_image_urls)
    """
    soup = BeautifulSoup(html, "html.parser")
    paragraphs = []
    image_urls = []

    for tag in soup.find_all(["p", "pre", "code"]):
        paragraphs.append(tag.get_text())

    for img in soup.find_all("img"):
        src = img.get("src")
        if src and src.startswith("http"):
            image_urls.append(src)

    full_text = "\n".join(paragraphs)
    return full_text, image_urls

def scrape_gfg_images(article_url, doc, save_folder="images"):
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"
    }

    # Create save folder if it doesn't exist
    os.makedirs(save_folder, exist_ok=True)

    # Step 1: Get HTML content
    response = requests.get(article_url, headers=headers)
    if response.status_code != 200:
        print("Failed to fetch page")
        return

    soup = BeautifulSoup(response.text, "html.parser")

    # Step 2: Extract all <img> tags
    img_tags = soup.find_all("img")

    print(f"Found {len(img_tags)} image(s)")

    allowed_exts = ['.png', '.jpg', '.jpeg']

    # Step 3: Download images
    for i, img_tag in enumerate(img_tags):
        src = img_tag.get("src") or img_tag.get("data-src")
        if not src:
            continue

        # Resolve relative URLs
        img_url = urljoin(article_url, src)

        # Check if the extension is allowed
        ext = os.path.splitext(img_url.split("?")[0])[1].lower()
        if ext not in allowed_exts:
            continue

        try:
            img_data = requests.get(img_url, headers=headers).content
            ext = img_url.split(".")[-1].split("?")[0][:4]  # Try to keep extension like png/jpg
            filename = f"gfg_img_{i+1}.{ext}"
            filepath = os.path.join(save_folder, filename)

            with open(filepath, "wb") as f:
                f.write(img_data)
            print(f"✅ Saved image: {filename}")

            if i == 0 or i == 1 or i == 2 or i == 3: 
                doc.add_picture(f"images/{filename}", width=Inches(5))
        except Exception as e:
            print(f"❌ Failed to download {img_url}: {e}")


def add_geeksforgeeks_content(topic, doc):
    """
    Uses Gemini to fetch GeeksforGeeks content on the topic and insert it into the document,
    cleaned of markdown and extraneous text.
    """
    prompt = (
        f"Extract content directly from the most relevant GeeksforGeeks about the topic '{topic}'. "
        f"Return only the actual content (no summaries, intros, or explanations about what you're doing). "
        f"Do not prepend or append anything. Preserve code snippets and examples as they appear. Give the full URL of the article as the first line"
    )
    gfg_raw = call_gemini_api(prompt)

    lines = gfg_raw.strip().split("\n", 1)
    # article_url = lines[0].strip()
    # article_url = "https://www.geeksforgeeks.org/agile-software-process-and-its-principles/"
    article_url = "https://www.geeksforgeeks.org/generative-adversarial-network-gan/"
    print("Article URL: ", article_url)
    scrape_gfg_images(article_url, doc)
    # doc.add_picture(f"images/gfg_img_1", width=Inches(5))

    gfg_raw = lines[1] if len(lines) > 1 else ""
    gfg_cleaned = clean_geeksforgeeks_content(gfg_raw)

    if gfg_cleaned:
        doc.add_heading("2. GeeksforGeeks Content", level=2)
        for paragraph in gfg_cleaned.split("\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
    else:
        doc.add_paragraph("Could not retrieve content from GeeksforGeeks.")


# def add_geeksforgeeks_content(topic, doc):
#     """
#     Uses Gemini to fetch GeeksforGeeks content on the topic and insert it into the document,
#     cleaned of markdown and extraneous text. Images from the content are downloaded and added too.
#     """
#     prompt = (
#         f"Extract raw HTML content directly from GeeksforGeeks about the topic '{topic}', including image tags. "
#         f"Return only the actual article content (no summaries, intros, or explanations about what you're doing). "
#         f"Do not prepend or append anything. Preserve all <img> tags and code snippets as they appear in the article."
#     )

#     gfg_raw_html = call_gemini_api(prompt)
#     print("GFG raw html", gfg_raw_html)

#     text, image_urls = extract_text_and_images(gfg_raw_html)
#     print("Image URLs", image_urls)

#     # Create image folder if not exists
#     image_folder = "images"
#     os.makedirs(image_folder, exist_ok=True)

#     for i, url in enumerate(image_urls):
#         try:
#             response = requests.get(url)
#             if (
#                 response.status_code == 200
#                 and "image" in response.headers["Content-Type"]
#             ):
#                 # Save image locally
#                 img_name = f"gfg_img_{i+1}.png"
#                 img_path = os.path.join(image_folder, img_name)

#                 img = Image.open(BytesIO(response.content)).convert("RGB")
#                 img.save(img_path, format="PNG")

#                 # Add to document
#                 doc.add_picture(img_path, width=Inches(5.5))
#                 doc.add_paragraph(f"Image source: {url}")
#             else:
#                 print(f"Failed to download image or content type not valid: {url}")
#         except Exception as e:
#             print(f"Error downloading image from {url}: {e}")

#         # if image_data:
#         #     doc.add_picture(image_data, width=Inches(5.5))
#         #     doc.add_paragraph(f"Image source: {url}")
#         # else:
#         #     doc.add_paragraph(f"[Image not added – download failed: {url}]")

#     if text:
#         doc.add_heading("2. GeeksforGeeks Content", level=2)
#         for paragraph in text.split("\n"):
#             if paragraph.strip():
#                 doc.add_paragraph(paragraph.strip())

#         for url in image_urls:
#             try:
#                 response = requests.get(url)
#                 image_stream = BytesIO(response.content)
#                 doc.add_picture(image_stream, width=Inches(5.5))
#             except Exception as e:
#                 doc.add_paragraph(f"[Image not loaded: {url}]")
#     else:
#         doc.add_paragraph("Could not retrieve content from GeeksforGeeks.")


def generate_summary(transcript, summary_length):
    """Generates a summary based on the selected length."""
    if summary_length == "Short":
        length_modifier = "30%"
    elif summary_length == "Medium":
        length_modifier = "50%"
    else:  # Detailed
        length_modifier = "keep all key information in detail."

    prompt = f"Refine and summarize the following transcript to {length_modifier} of the original transcript and give it a suitable topic name and write the content in the form 'Topic: followed by the rest of the content', also don't make any text bold:\n{transcript}"
    refined_summary = call_gemini_api(prompt)
    # print("Refined summary: ", refined_summary)
    return (
        refined_summary if refined_summary else transcript
    )  # Fallback to original transcript


def transcript_to_json(transcript):
    """Converts the refined transcript into a structured JSON format for slide creation."""
    slides = []
    # Extract the topic from the beginning of the transcript
    topic_match = re.search(r"Topic: (.+)", transcript)
    if topic_match:
        title = topic_match.group(1).strip()
        # Remove the "Topic: ..." part from the transcript content
        content = transcript.replace(topic_match.group(0), "").strip()
        sentences = sent_tokenize(content)
        slides.append({"title": title, "points": sentences})
    else:
        # If no "Topic: ..." is found, use the entire transcript as content
        sentences = sent_tokenize(transcript)
        slides.append(
            {"title": "Transcript", "points": sentences}
        )  # or some default title.
    return slides


def generate_notes(transcript, summary_length):
    """Generates notes, refines them with Gemini, and saves to a Word document."""

    # 1. Transcript Summary
    refined_transcript = generate_summary(transcript, summary_length)

    refined_transcript_json = transcript_to_json(refined_transcript)

    # print("Refined Transcript: ", refined_transcript)

    with open("slides.json", "w") as f:
        json.dump(refined_transcript_json, f, indent=4)
    with open("slides.json", "r") as file:
        slides_data = json.load(file)
    topic = slides_data[0]["title"]
    # print("TOPIC::", topic)
    doc = Document()
    add_page_border(doc)  # Add border to pages
    add_headers_and_footers(doc, topic)  # Add header and footer

    doc.add_heading(f"Lecture Notes on {topic}", level=1)
    doc.add_heading("1. Transcript Summary", level=2)

    # print(topic)
    identify_and_bold_subtopics(
        (
            clean_text(refined_transcript)
            if refined_transcript
            else clean_text(transcript)
        ),
        doc,
    )
    add_geeksforgeeks_content(topic, doc)
    # def add_geeksforgeeks_content(topic, doc):
    #     """Uses Gemini to fetch GeeksforGeeks content for the topic and insert it into the document without modifying it."""
    #     prompt = (
    #         f"Search GeeksforGeeks for an article on the topic '{topic}' and extract the most relevant content. "
    #         f"Return only the exact content from the article without modifying any word. Do not summarize or refine. "
    #         f"Only include educational content, not website elements or advertisements."
    #     )
    #     gfg_content = call_gemini_api(prompt)

    #     if gfg_content:
    #         doc.add_heading("3. GeeksforGeeks Content", level=2)
    #         for paragraph in gfg_content.split("\n"):
    #             if paragraph.strip():
    #                 doc.add_paragraph(paragraph.strip())
    #     else:
    #         doc.add_paragraph("Could not retrieve content from GeeksforGeeks.")

    # 3. Additional Web Content
    doc.add_heading("3. Additional Web Content", level=2)
    related_links = search_google_free(topic, num_results=3)
    for link in related_links:
        web_content = scrape_web_content(link)
        refined_web_content = call_gemini_api(
            f"Refine and summarize the following article and add some additional information:\n{web_content}"
        )
        identify_and_bold_subtopics(
            (
                clean_text(refined_web_content)
                if refined_web_content
                else clean_text(web_content)
            ),
            doc,
        )

    # 4. Key Takeaways
    doc.add_heading("4. Key Takeaways", level=2)
    key_points = [
        "Introduction to the topic",
        "Key discoveries and developments",
        "Future advancements and applications",
    ]
    refined_key_points = call_gemini_api(
        f"Generate detailed bullet points for the following key areas related to {topic}:\n{key_points}"
    )
    identify_and_bold_subtopics(
        clean_text(refined_key_points) if refined_key_points else "\n".join(key_points),
        doc,
    )

    # Sources
    doc.add_heading("5. Sources", level=2)
    for source in related_links:
        doc.add_paragraph(f"- {source}")

    filename = "Detailed_Lecture_Notes.docx"
    doc.save(filename)

    # Save document in memory
    notes_stream = BytesIO()
    doc.save(notes_stream)
    notes_stream.seek(0)

    return notes_stream

    # print(f"✅ Notes saved as {filename}")


# if __name__ == "__main__":
#     sample_transcript = """Artificial intelligence is transforming industries worldwide.
#     Machine learning models improve decision-making, automate tasks, and enhance efficiency.
#     However, AI ethics and biases remain crucial concerns.
#     The future of AI will be shaped by policy decisions and technological advancements."""

#     print("Refined Transcript:", call_gemini_api(sample_transcript))
#     generate_notes(sample_transcript, "Artificial Intelligence")
